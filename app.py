import os
import streamlit as st
from dotenv import load_dotenv
from google import generativeai as genai
import pdfplumber
import docx

load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

genai.configure()
model = genai.GenerativeModel('gemini-1.5-flash')


st.set_page_config(page_title="Job Application Booster", page_icon="💼")

st.title("💼 Job Application Booster")
st.subheader("Step 2: Design Your Application Assistant Interface")


st.header("1. Upload Your Files")
# Resume
uploaded_file = st.file_uploader("Upload Resume (.pdf or .docx)", type=["pdf", "docx"])

extracted_resume = ""
if uploaded_file:
    file_type = uploaded_file.name.split('.')[-1]
    try:
        if file_type == "pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                extracted_resume = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif file_type == "docx":
            doc = docx.Document(uploaded_file)
            extracted_resume = "\n".join([para.text for para in doc.paragraphs])
        else:
            st.warning("Unsupported file type.")
    except Exception as e:
        st.error(f"Error extracting text: {e}")

resume = st.text_area("📄 Paste your Resume here", height=250, value=extracted_resume, help="Copy and paste your resume.")


# JD
uploaded_file = st.file_uploader("Upload Job Description (.pdf or .docx)", type=["pdf", "docx"])
extracted_jd = ""

if uploaded_file:
    file_type = uploaded_file.name.split('.')[-1]
    try:
        if file_type == "pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                extracted_jd = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif file_type == "docx":
            doc = docx.Document(uploaded_file)
            extracted_jd = "\n".join([para.text for para in doc.paragraphs])
        else:
            st.warning("Unsupported file type.")
    except Exception as e:
        st.error(f"Error extracting text: {e}")


job_desc = st.text_area("📌 Paste the Job Description here", height=250, value=extracted_jd, help="Copy and paste the full job post.")

st.slider("🎛️ Temperature (Model Creativity))", 0.0, 1.0, 0.7, 0.1, key="temperature")

st.header("2. Choose What to Generate")

col1, col2, col3 = st.columns(3)
with col1:
    generate_cl = st.checkbox("✉️ Cover Letter", value=True)
with col2:
    generate_ri = st.checkbox("📝 Resume Improvements", value=True)
with col3:
    generate_it = st.checkbox("💡 Interview Tips", value=True)

st.header("3. Boost!")

if st.button("🚀 Boost My Application"):
    if not resume or not job_desc:
        st.warning("⚠️ Please fill out both Resume and Job Description.")
    elif not (generate_cl or generate_ri or generate_it):
        st.warning("⚠️ Please select at least one generation option.")
    else:
        with st.spinner("✨ Gemini is generating..."):
            items = []
            if generate_cl:
                items.append("### Cover Letter\nWrite a tailored cover letter to 'Hiring Manager'.")
            if generate_ri:
                items.append("### Resume Improvements\nList 3-5 actionable resume improvements, focusing on keywords and relevance to the job description.")
            if generate_it:
                items.append("### Interview Tips\nList 3-5 concise interview preparation tips specific to this role.")

            full_prompt = f"""
            You are a professional career coach and an expert in job applications.
            Given the following resume and job description, please generate the requested items in clearly separated sections using headings (###).
            Ensure the outputs are professional, concise, and highly relevant
            
            Resume: {resume}
            Job Description: {job_desc}
            Please provide: {" ".join(items)}
            """

            try:
                # response = model.generate_content(full_prompt)
                # response_text = response.text
                response_stream = model.generate_content(full_prompt, 
                                                         stream=True, 
                                                         generation_config={"temperature": st.session_state.temperature}
                                                        )

                full_response = ""
                response_placeholder = st.empty()
                for chunk in response_stream:
                    full_response += chunk.text
                    response_placeholder.markdown(full_response + "▌")

                # sections = {"Cover Letter": "", "Resume Improvements": "", "Interview Tips": ""}
                # current_section = None
    
                # for line in response_text.splitlines():
                #     if line.startswith("###"):
                #         current_section = line.replace("###", "").strip()
                #         sections[current_section] = ""
                #     elif current_section:
                #         sections[current_section] += line + "\n"

                # st.markdown("---")
                # if generate_cl:
                #     st.subheader("✉️ Cover Letter")
                #     st.markdown(sections["Cover Letter"])
                #     st.code(sections["Cover Letter"], language="text")
    
                # if generate_ri:
                #     st.subheader("📝 Resume Improvements")
                #     st.markdown(sections["Resume Improvements"])
                #     st.code(sections["Resume Improvements"], language="text")
    
                # if generate_it:
                #     st.subheader("💡 Interview Tips")
                #     st.markdown(sections["Interview Tips"])
                #     st.code(sections["Interview Tips"], language="text")
            except Exception as e:
                st.error(f"Gemini API Error: {e}")


st.markdown("---")
st.caption("Built with ❤️ by QubyteFlow")